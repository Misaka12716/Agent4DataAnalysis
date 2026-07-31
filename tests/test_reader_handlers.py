"""Reader Handler 对各白名单格式的解析覆盖测试。"""

from __future__ import annotations

import json
from pathlib import Path

import pandas as pd
import pytest
from PIL import Image

from reader.handlers.document_docx import digest_docx_file
from reader.handlers.document_pdf import digest_pdf_file
from reader.handlers.image import digest_image_file
from reader.handlers.table import _read_raw_table, digest_table_file
from reader.handlers.text import digest_text_file


def _assert_no_missing_lib_note(entry: dict) -> None:
    note = entry.get("note") or ""
    assert "未安装" not in note, f"unexpected missing-lib note: {note}"


def test_digest_csv_and_tsv(tmp_path: Path):
    csv_path = tmp_path / "a.csv"
    csv_path.write_text("name,age\nAlice,30\n", encoding="utf-8")
    tsv_path = tmp_path / "b.tsv"
    tsv_path.write_text("name\tage\nBob\t20\n", encoding="utf-8")

    csv_entry = digest_table_file(str(tmp_path), "a.csv")
    tsv_entry = digest_table_file(str(tmp_path), "b.tsv")
    _assert_no_missing_lib_note(csv_entry)
    _assert_no_missing_lib_note(tsv_entry)
    assert csv_entry.get("error") is None
    assert "name" in csv_entry.get("columns", [])
    assert tsv_entry.get("error") is None
    assert "name" in tsv_entry.get("columns", [])


def test_digest_csv_gbk_encoding(tmp_path: Path):
    fp = tmp_path / "cn.csv"
    fp.write_bytes("姓名,年龄\n张三,18\n".encode("gbk"))
    entry = digest_table_file(str(tmp_path), "cn.csv")
    _assert_no_missing_lib_note(entry)
    assert entry.get("error") is None
    assert any("姓名" in str(c) or "张" in str(c) for c in entry.get("columns", [])) or entry.get(
        "shape", [0]
    )[0] >= 1


def test_digest_xlsx(tmp_path: Path):
    fp = tmp_path / "data.xlsx"
    pd.DataFrame({"a": [1, 2], "b": [3, 4]}).to_excel(fp, index=False, engine="openpyxl")
    entry = digest_table_file(str(tmp_path), "data.xlsx")
    _assert_no_missing_lib_note(entry)
    assert entry.get("error") is None
    assert "a" in entry.get("columns", [])
    assert entry.get("shape", [0, 0])[0] == 2


def test_read_raw_table_xlsx_uses_openpyxl(tmp_path: Path):
    fp = tmp_path / "t.xlsx"
    pd.DataFrame({"x": [1]}).to_excel(fp, index=False, engine="openpyxl")
    df = _read_raw_table(str(fp), ".xlsx")
    assert df.shape[0] >= 1


@pytest.mark.skipif(
    __import__("importlib").util.find_spec("xlrd") is None,
    reason="xlrd not installed",
)
def test_digest_xls_with_xlrd(tmp_path: Path):
    fp = tmp_path / "legacy.xls"
    try:
        pd.DataFrame({"a": [1, 2]}).to_excel(fp, index=False, engine="xlwt")
    except Exception:
        pytest.skip("xlwt not available to create .xls fixture")
    entry = digest_table_file(str(tmp_path), "legacy.xls")
    _assert_no_missing_lib_note(entry)
    assert entry.get("error") is None


def test_digest_docx_paragraphs_and_tables(tmp_path: Path):
    docx = pytest.importorskip("docx")
    from docx import Document

    fp = tmp_path / "note.docx"
    doc = Document()
    doc.add_paragraph("标题段落")
    doc.add_paragraph("正文内容")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "列A"
    table.cell(0, 1).text = "列B"
    table.cell(1, 0).text = "1"
    table.cell(1, 1).text = "2"
    doc.save(str(fp))

    entry = digest_docx_file(str(tmp_path), "note.docx")
    _assert_no_missing_lib_note(entry)
    assert entry.get("error") is None
    assert entry.get("paragraph_count", 0) >= 2
    assert entry.get("table_count", 0) >= 1
    preview = entry.get("preview") or ""
    assert "标题段落" in preview
    assert "列A" in preview


def test_digest_pdf(tmp_path: Path):
    # 优先用 reportlab 生成；不可用则用最小 PDF 字节
    fp = tmp_path / "doc.pdf"
    try:
        from reportlab.pdfgen import canvas

        c = canvas.Canvas(str(fp))
        c.drawString(100, 750, "Hello PDF Preview")
        c.save()
    except ImportError:
        # 最小合法 PDF（无文本层时 preview 可能为空，但仍应成功打开）
        fp.write_bytes(
            b"""%PDF-1.1
1 0 obj<< /Type /Catalog /Pages 2 0 R >>endobj
2 0 obj<< /Type /Pages /Kids [3 0 R] /Count 1 >>endobj
3 0 obj<< /Type /Page /Parent 2 0 R /MediaBox [0 0 300 144] /Contents 4 0 R /Resources<< /Font<< /F1 5 0 R >> >> >>endobj
4 0 obj<< /Length 44 >>stream
BT /F1 12 Tf 100 100 Td (Hello PDF) Tj ET
endstream
endobj
5 0 obj<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>endobj
xref
0 6
0000000000 65535 f 
0000000009 00000 n 
0000000058 00000 n 
0000000115 00000 n 
0000000266 00000 n 
0000000361 00000 n 
trailer<< /Size 6 /Root 1 0 R >>
startxref
429
%%EOF"""
        )

    entry = digest_pdf_file(str(tmp_path), "doc.pdf")
    _assert_no_missing_lib_note(entry)
    assert entry.get("error") is None
    assert entry.get("page_count", 0) >= 1
    assert entry.get("parser") in ("pypdf", "PyPDF2")


def test_digest_text_and_json(tmp_path: Path):
    (tmp_path / "a.txt").write_text("hello\nworld\n", encoding="utf-8")
    (tmp_path / "b.json").write_text(json.dumps({"foo": 1, "bar": [1, 2]}), encoding="utf-8")

    text_entry = digest_text_file(str(tmp_path), "a.txt")
    json_entry = digest_text_file(str(tmp_path), "b.json")
    assert "hello" in (text_entry.get("preview") or "")
    assert json_entry.get("json_type") == "dict"
    assert "foo" in (json_entry.get("json_keys") or [])


def test_digest_png(tmp_path: Path):
    fp = tmp_path / "img.png"
    Image.new("RGB", (16, 8), color=(255, 0, 0)).save(fp)
    entry = digest_image_file(str(tmp_path), "img.png")
    assert entry.get("error") is None
    assert entry.get("width") == 16
    assert entry.get("height") == 8


def test_check_reader_parse_deps_runs():
    from reader.deps import check_reader_parse_deps, reset_deps_check_for_tests

    reset_deps_check_for_tests()
    status = check_reader_parse_deps(force=True)
    assert isinstance(status, dict)
    assert "python-docx" in status
    assert "Pillow" in status
