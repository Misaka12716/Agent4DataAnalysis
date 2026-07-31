"""DOCX 文档 Handler（依赖 python-docx，缺失时降级）。"""

from __future__ import annotations

import os
from typing import Any, Dict, List


def _table_to_text(table: Any) -> str:
    rows: List[str] = []
    for row in table.rows:
        cells = [(c.text or "").strip().replace("\n", " ") for c in row.cells]
        if any(cells):
            rows.append(" | ".join(cells))
    return "\n".join(rows)


def digest_docx_file(workspace_root: str, relative_path: str, **_kwargs: Any) -> Dict[str, Any]:
    fp = os.path.join(workspace_root, relative_path.replace("/", os.sep))
    entry: Dict[str, Any] = {
        "file_type": "document",
        "format": "docx",
        "relative_path": relative_path,
        "handler_id": "document_docx",
    }
    try:
        entry["file_size_bytes"] = os.path.getsize(fp)
    except OSError as e:
        entry["error"] = str(e)
        return entry

    try:
        from docx import Document  # type: ignore
    except ImportError:
        entry["note"] = "未安装 python-docx，仅登记元数据；可 pip install python-docx 启用文本抽取"
        return entry

    try:
        doc = Document(fp)
        paras = [p.text.strip() for p in doc.paragraphs if (p.text or "").strip()]
        table_texts: List[str] = []
        for i, table in enumerate(doc.tables):
            t = _table_to_text(table)
            if t:
                table_texts.append(f"[表{i + 1}]\n{t}")

        parts: List[str] = []
        if paras:
            parts.append("\n".join(paras))
        if table_texts:
            parts.append("\n\n".join(table_texts))
        full = "\n\n".join(parts)
        preview = full[:2000]
        entry.update(
            {
                "paragraph_count": len(paras),
                "table_count": len(doc.tables),
                "preview": preview,
                "preview_truncated": len(full) > 2000,
                "parser": "python-docx",
            }
        )
    except Exception as e:
        entry["error"] = str(e)
    return entry


class DocxDocumentHandler:
    handler_id = "document_docx"

    def digest(self, workspace_root: str, relative_path: str, **kwargs: Any) -> Dict[str, Any]:
        return digest_docx_file(workspace_root, relative_path, **kwargs)
